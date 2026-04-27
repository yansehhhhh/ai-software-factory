#!/usr/bin/env python3
"""
将 Markdown 文件转换为 Word 文档 (.docx)
用法: python md2docx.py <input.md> <output.docx>
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)


def convert_md_to_docx(md_path: str, docx_path: str):
    """将 Markdown 文件转换为 Word 文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 设置标题样式
    for level in range(1, 7):
        heading_style = doc.styles[f"Heading {level}"]
        heading_font = heading_style.font
        heading_font.name = "微软雅黑"
        heading_font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        heading_style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        if level == 1:
            heading_font.size = Pt(22)
            heading_font.bold = True
        elif level == 2:
            heading_font.size = Pt(18)
            heading_font.bold = True
        elif level == 3:
            heading_font.size = Pt(15)
            heading_font.bold = True
        elif level == 4:
            heading_font.size = Pt(14)
            heading_font.bold = True
        elif level == 5:
            heading_font.size = Pt(13)
            heading_font.bold = True
        elif level == 6:
            heading_font.size = Pt(12)
            heading_font.bold = True

    content = Path(md_path).read_text(encoding="utf-8")
    lines = content.split("\n")

    in_code_block = False
    in_table = False
    table_rows = []
    table_header_done = False
    code_paragraph = None

    i = 0
    while i < len(lines):
        line = lines[i]

        # 代码块处理
        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                if code_paragraph:
                    code_paragraph.style = doc.styles["Normal"]
                    # 设置代码块样式
                    for run in code_paragraph.runs:
                        run.font.name = "Consolas"
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    # 添加灰色背景（通过 shading）
                    shading_elm = code_paragraph._p.get_or_add_pPr().makeelement(
                        qn("w:shd"),
                        {
                            qn("w:fill"): "F5F5F5",
                            qn("w:val"): "clear",
                        },
                    )
                    code_paragraph._p.get_or_add_pPr().append(shading_elm)
                code_paragraph = None
            else:
                in_code_block = True
                code_paragraph = doc.add_paragraph()
                code_paragraph.style = doc.styles["Normal"]
            i += 1
            continue

        if in_code_block:
            if code_paragraph is not None:
                if code_paragraph.text:
                    code_paragraph.add_run("\n")
                run = code_paragraph.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # 表格处理
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_rows = []

            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            # 检查是否是分隔行
            if all(set(c.strip().replace("-", "")) == {"-"} for c in cells if c.strip()):
                table_header_done = True
                i += 1
                continue

            table_rows.append(cells)
            i += 1

            # 查看下一行是否还是表格
            if i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                continue

            # 表格结束，创建 Word 表格
            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = "Medium Shading 1 Accent 1"

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_text
                        # 设置单元格字体
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = "微软雅黑"
                                run.font.size = Pt(10)
                                paragraph.style.element.rPr.rFonts.set(
                                    qn("w:eastAsia"), "微软雅黑"
                                )
                doc.add_paragraph()  # 表格后空一行

            in_table = False
            table_rows = []
            table_header_done = False
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 标题
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # 无序列表
        if re.match(r"^\s*[-*+]\s+", line):
            text = re.sub(r"^\s*[-*+]\s+", "", line).strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
            i += 1
            continue

        # 有序列表
        ordered_match = re.match(r"^\s*\d+\.\s+(.*)", line)
        if ordered_match:
            text = ordered_match.group(1).strip()
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_formatted_text(p, line.strip())
        i += 1

    doc.save(docx_path)
    print(f"已生成: {docx_path}")


def _add_formatted_text(paragraph, text):
    """解析 Markdown 内联格式（粗体、斜体、行内代码）并添加到段落"""
    # 处理行内代码
    parts = re.split(r"`([^`]+)`", text)
    for idx, part in enumerate(parts):
        if idx % 2 == 1:
            # 行内代码
            run = paragraph.add_run(part)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xE8, 0x3E, 0x80)
            # 灰色背景
            shading_elm = run._r.get_or_add_rPr().makeelement(
                qn("w:shd"),
                {
                    qn("w:fill"): "F0F0F0",
                    qn("w:val"): "clear",
                },
            )
            run._r.get_or_add_rPr().append(shading_elm)
        else:
            # 处理粗体和斜体
            _add_bold_italic(paragraph, part)


def _add_bold_italic(paragraph, text):
    """处理粗体和斜体"""
    # ***粗斜体***
    parts = re.split(r"\*\*\*([^*]+)\*\*\*", text)
    result_runs = []

    for idx, part in enumerate(parts):
        if idx % 2 == 1:
            result_runs.append(("bold_italic", part))
        else:
            # **粗体**
            sub_parts = re.split(r"\*\*([^*]+)\*\*", part)
            for sidx, spart in enumerate(sub_parts):
                if sidx % 2 == 1:
                    result_runs.append(("bold", spart))
                else:
                    # *斜体*
                    inner_parts = re.split(r"\*([^*]+)\*", spart)
                    for iidx, ipart in enumerate(inner_parts):
                        if iidx % 2 == 1:
                            result_runs.append(("italic", ipart))
                        else:
                            result_runs.append(("normal", ipart))

    for style, text in result_runs:
        if not text:
            continue
        run = paragraph.add_run(text)
        if style == "bold":
            run.bold = True
        elif style == "italic":
            run.italic = True
        elif style == "bold_italic":
            run.bold = True
            run.italic = True


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python md2docx.py <input.md> <output.docx>")
        sys.exit(1)

    convert_md_to_docx(sys.argv[1], sys.argv[2])